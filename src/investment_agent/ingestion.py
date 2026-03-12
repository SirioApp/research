from __future__ import annotations

import json
import re
import time
from pathlib import Path
from urllib.parse import urlparse

from .models import SourceDocument


class IngestionError(RuntimeError):
    pass


class SourceIngestor:
    _HTTP_HEADERS = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/126.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/json;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Cache-Control": "no-cache",
        "Pragma": "no-cache",
    }

    def ingest(self, sources: list[str]) -> list[SourceDocument]:
        documents: list[SourceDocument] = []
        for source in sources:
            if self._is_url(source):
                text = self._read_url(source)
            else:
                text = self._read_file(Path(source))
            normalized = self._normalize_whitespace(text)
            if normalized:
                documents.append(SourceDocument(source=source, text=normalized))

        if not documents:
            raise IngestionError("No readable content found in the provided sources.")

        return documents

    @staticmethod
    def _is_url(value: str) -> bool:
        parsed = urlparse(value)
        return parsed.scheme in {"http", "https"} and bool(parsed.netloc)

    def _read_url(self, url: str) -> str:
        try:
            import requests
            from bs4 import BeautifulSoup
        except ImportError as exc:
            raise IngestionError(
                "Missing dependencies for URL ingestion. Install requests and beautifulsoup4."
            ) from exc

        session = requests.Session()
        session.headers.update(self._HTTP_HEADERS)

        response = self._request_with_retry(session, url)
        if response is None:
            fallback_text = self._read_via_jina_fallback(session, url)
            if fallback_text:
                return fallback_text
            raise IngestionError(f"Unable to fetch URL after retries (possibly rate limited): {url}")

        content_type = response.headers.get("content-type", "")
        if "application/json" in content_type:
            payload = response.json()
            return json.dumps(payload, ensure_ascii=True, indent=2)

        html = response.text
        soup = BeautifulSoup(html, "html.parser")
        for node in soup(["script", "style", "noscript"]):
            node.decompose()
        return soup.get_text(" ")

    def _request_with_retry(self, session, url: str, max_attempts: int = 5):
        retryable_status = {429, 500, 502, 503, 504}

        for attempt in range(1, max_attempts + 1):
            try:
                response = session.get(url, timeout=25, allow_redirects=True)
            except Exception:
                response = None

            if response is not None and response.status_code < 400:
                return response

            if response is not None and response.status_code not in retryable_status:
                if response.status_code == 403:
                    return None
                response.raise_for_status()

            if attempt == max_attempts:
                return None

            retry_after = None
            if response is not None:
                retry_after = self._parse_retry_after(response.headers.get("Retry-After"))
            time.sleep(self._backoff_seconds(attempt, retry_after))

        return None

    def _read_via_jina_fallback(self, session, url: str) -> str | None:
        normalized = url.removeprefix("https://").removeprefix("http://")
        fallback_url = f"https://r.jina.ai/http://{normalized}"
        response = self._request_with_retry(session, fallback_url, max_attempts=3)
        if response is None or response.status_code >= 400:
            return None
        return response.text

    @staticmethod
    def _parse_retry_after(value: str | None) -> float | None:
        if value is None:
            return None
        try:
            return max(0.0, float(value))
        except ValueError:
            return None

    @staticmethod
    def _backoff_seconds(attempt: int, retry_after: float | None) -> float:
        if retry_after is not None:
            return min(retry_after, 20.0)
        return min(20.0, (2 ** (attempt - 1)) * 1.2)

    def _read_file(self, path: Path) -> str:
        if not path.exists() or not path.is_file():
            raise IngestionError(f"Source file not found: {path}")

        extension = path.suffix.lower()
        if extension in {".txt", ".md", ".rst", ".log", ".csv"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if extension == ".json":
            raw = path.read_text(encoding="utf-8", errors="ignore")
            return json.dumps(json.loads(raw), ensure_ascii=True, indent=2)
        if extension == ".pdf":
            return self._read_pdf(path)
        if extension == ".docx":
            return self._read_docx(path)

        return path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _read_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise IngestionError("Missing pypdf dependency for PDF ingestion.") from exc

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _read_docx(path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise IngestionError("Missing python-docx dependency for DOCX ingestion.") from exc

        document = Document(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        collapsed = re.sub(r"\s+", " ", text).strip()
        return collapsed
